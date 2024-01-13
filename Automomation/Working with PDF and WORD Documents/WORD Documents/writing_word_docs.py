#! python 3

import docx

doc_file = docx.Document()
doc_file.add_paragraph('Hello World!')

doc_paragraph = doc_file.add_paragraph('This is a second paragraph')
doc_paragraph = doc_file.add_paragraph('This is yet another paragraph.')
doc_paragraph.add_run('This text is being added to the second paragraph')
doc_file.save('multiple_paragraph.docx')
