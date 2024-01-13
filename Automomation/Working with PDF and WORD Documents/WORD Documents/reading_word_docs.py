#! python 3

import docx

word_doc = "C:\\Users\\User\Desktop\Functional Safety\FTA\FTA Basics.docx"

# Read the document file
doc_file = docx.Document(word_doc)

# Finding the number of paragraphs in the document
doc_length = len(doc_file.paragraphs)

# Printing each paragraph in the document
for paragraph in range(doc_length):
    print(doc_file.paragraphs[paragraph].text)

# Printing the runs from each paragraph
runs_length_list = []

for count in range(doc_length):
    runs_length_list.append(len(doc_file.paragraphs[count].runs))

total_runs = sum(runs_length_list)

print(runs_length_list)
print(total_runs)

#
print(doc_file.paragraphs[1].runs[0].text)
print(doc_file.paragraphs[1].runs[1].text)
